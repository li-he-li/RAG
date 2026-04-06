"""
API router for similarity search endpoints.
"""

from __future__ import annotations

import io
import logging
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import JSONResponse, StreamingResponse
from sqlalchemy import func
from sqlalchemy.orm import Session

from app.core.config import get_bootstrap_status, get_bootstrap_missing_components
from app.core.database import get_session
from app.models.db_tables import DocumentTable, ParagraphTable
from app.models.schemas import (
    BootstrapStatus,
    ChatRequest,
    ChatResponse,
    ContractReviewRequest,
    DocumentIngestRequest,
    DocumentIngestResponse,
    DocumentListItem,
    ErrorResponse,
    ReviewTemplateRecommendationResponse,
    SimilarCaseSearchRequest,
    SimilarCaseSearchResponse,
    SessionTempClearResponse,
    SessionTempFileItem,
    SessionTempFileKind,
    SearchRequest,
    SearchResponse,
)
from app.services.indexer import ingest_document, delete_document
from app.services.contract_review import stream_template_difference_review
from app.services.similar_case_search import execute_similar_case_search
from app.services.session_files import session_temp_file_store
from app.services.template_recommendation import recommend_templates_for_session
from app.services.chat import execute_grounded_chat, stream_grounded_chat
from app.services.retrieval import execute_search
from app.services.traceability import validate_and_enrich_results, TraceabilityValidationError

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api", tags=["legal-search"])
DOC_PREVIEW_TEXT_RE = re.compile(r"[\u4e00-\u9fffA-Za-z0-9][^\x00]{3,}")


def _ensure_retrieval_ready() -> None:
    """Preflight gate: block retrieval endpoints until bootstrap completes."""
    status = get_bootstrap_status()
    if status.get("all_ready", False):
        return
    missing = get_bootstrap_missing_components()
    detail = "Retrieval service is unavailable until bootstrap completes."
    if missing:
        detail = f"{detail} Missing: {', '.join(missing)}"
    raise HTTPException(status_code=503, detail=detail)


def _extract_upload_text(file_name: str, raw: bytes) -> str:
    """Extract plain text from supported upload types."""
    ext = Path(file_name).suffix.lower()

    if ext in {".txt", ".md"}:
        return raw.decode("utf-8", errors="replace")

    if ext == ".pdf":
        try:
            from pypdf import PdfReader
        except Exception as exc:
            raise HTTPException(
                status_code=500,
                detail=f"PDF 解析依赖未安装: {exc}",
            ) from exc

        reader = PdfReader(io.BytesIO(raw))
        text = "\n".join((page.extract_text() or "").strip() for page in reader.pages)
        if not text.strip():
            raise HTTPException(status_code=422, detail="PDF 未提取到可用文本内容")
        return text

    if ext == ".docx":
        try:
            from docx import Document
        except Exception as exc:
            raise HTTPException(
                status_code=500,
                detail=f"DOCX 解析依赖未安装: {exc}",
            ) from exc

        doc = Document(io.BytesIO(raw))
        text = "\n".join((p.text or "").strip() for p in doc.paragraphs)
        if not text.strip():
            raise HTTPException(status_code=422, detail="DOCX 未提取到可用文本内容")
        return text

    if ext == ".doc":
        raise HTTPException(
            status_code=415,
            detail="暂不支持 .doc，请先另存为 .docx 后上传",
        )

    raise HTTPException(
        status_code=415,
        detail="不支持的文件类型，仅支持 .txt .md .pdf .docx",
    )

def _normalize_extracted_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    cleaned: list[str] = []
    previous = None
    for line in lines:
        compact = re.sub(r"\s+", " ", line).strip()
        if not compact or compact == previous:
            continue
        cleaned.append(compact)
        previous = compact
    return "\n".join(cleaned).strip()


def _ensure_non_empty_text(text: str, detail: str) -> str:
    normalized = _normalize_extracted_text(text)
    if not normalized:
        raise HTTPException(status_code=422, detail=detail)
    return normalized


def _extract_xlsx_text(raw: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"XLSX parser unavailable: {exc}") from exc

    workbook = load_workbook(filename=io.BytesIO(raw), data_only=True, read_only=True)
    rows: list[str] = []
    for sheet in workbook.worksheets:
        rows.append(f"[工作表] {sheet.title}")
        for values in sheet.iter_rows(values_only=True):
            cells = [str(value).strip() for value in values if value is not None and str(value).strip()]
            if cells:
                rows.append(" | ".join(cells))
    return _ensure_non_empty_text("\n".join(rows), "Excel 文件未提取到可用文本内容")


def _extract_xls_text(raw: bytes) -> str:
    try:
        import xlrd
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"XLS parser unavailable: {exc}") from exc

    workbook = xlrd.open_workbook(file_contents=raw)
    rows: list[str] = []
    for sheet in workbook.sheets():
        rows.append(f"[工作表] {sheet.name}")
        for row_idx in range(sheet.nrows):
            values: list[str] = []
            for col_idx in range(sheet.ncols):
                value = sheet.cell_value(row_idx, col_idx)
                if value is None:
                    continue
                text = str(value).strip()
                if text:
                    values.append(text)
            if values:
                rows.append(" | ".join(values))
    return _ensure_non_empty_text("\n".join(rows), "Excel 文件未提取到可用文本内容")


def _extract_doc_text_via_word(raw: bytes) -> str:
    try:
        import pythoncom
        import win32com.client
    except Exception:
        return ""

    with tempfile.TemporaryDirectory() as temp_dir:
        source_path = Path(temp_dir) / "upload.doc"
        output_path = Path(temp_dir) / "upload.txt"
        source_path.write_bytes(raw)

        pythoncom.CoInitialize()
        word = None
        doc = None
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            doc = word.Documents.Open(str(source_path), ReadOnly=True)
            doc.SaveAs(str(output_path), FileFormat=7)
        finally:
            if doc is not None:
                doc.Close(False)
            if word is not None:
                word.Quit()
            pythoncom.CoUninitialize()

        if not output_path.exists():
            return ""
        return output_path.read_text(encoding="utf-16", errors="ignore")


def _extract_doc_text_via_antiword(raw: bytes) -> str:
    antiword_path = shutil.which("antiword")
    if not antiword_path:
        return ""

    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as temp_file:
        temp_path = Path(temp_file.name)
        temp_file.write(raw)

    try:
        completed = subprocess.run(
            [antiword_path, str(temp_path)],
            check=False,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="ignore",
        )
        if completed.returncode != 0:
            return ""
        return completed.stdout
    finally:
        temp_path.unlink(missing_ok=True)


def _extract_doc_stream_candidates(stream: bytes) -> list[str]:
    candidates: list[str] = []
    for encoding, offsets in (("utf-16le", (0, 1)), ("gb18030", (0,)), ("utf-8", (0,)), ("latin1", (0,))):
        for offset in offsets:
            if offset >= len(stream):
                continue
            try:
                decoded = stream[offset:].decode(encoding, errors="ignore")
            except Exception:
                continue
            decoded = decoded.replace("\x00", " ")
            matches = DOC_PREVIEW_TEXT_RE.findall(decoded)
            if matches:
                candidates.append("\n".join(matches))
    return candidates


def _extract_doc_text_via_ole(raw: bytes) -> str:
    try:
        import olefile
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"DOC parser unavailable: {exc}") from exc

    candidates: list[str] = []
    with olefile.OleFileIO(io.BytesIO(raw)) as ole:
        for stream_name in ole.listdir(streams=True, storages=False):
            try:
                stream = ole.openstream(stream_name).read()
            except Exception:
                continue
            candidates.extend(_extract_doc_stream_candidates(stream))

    unique_lines: list[str] = []
    seen = set()
    for candidate in candidates:
        for line in candidate.splitlines():
            compact = re.sub(r"\s+", " ", line).strip()
            if len(compact) < 4 or compact in seen:
                continue
            seen.add(compact)
            unique_lines.append(compact)
    return "\n".join(unique_lines)


def _extract_doc_text(raw: bytes) -> str:
    for extractor in (_extract_doc_text_via_word, _extract_doc_text_via_antiword, _extract_doc_text_via_ole):
        try:
            text = extractor(raw)
        except HTTPException:
            raise
        except Exception:
            logger.debug("DOC extractor failed: %s", extractor.__name__, exc_info=True)
            continue
        if text:
            return _ensure_non_empty_text(text, "DOC 文件未提取到可用文本内容")

    raise HTTPException(status_code=422, detail="DOC 文件未提取到可用文本内容")


def _extract_upload_text(file_name: str, raw: bytes) -> str:
    """Extract plain text from supported upload types."""
    ext = Path(file_name).suffix.lower()

    if ext in {".txt", ".md"}:
        return _ensure_non_empty_text(raw.decode("utf-8", errors="replace"), "文件未提取到可用文本内容")

    if ext == ".pdf":
        try:
            from pypdf import PdfReader
        except Exception as exc:
            raise HTTPException(status_code=500, detail=f"PDF parser unavailable: {exc}") from exc

        reader = PdfReader(io.BytesIO(raw))
        text = "\n".join((page.extract_text() or "").strip() for page in reader.pages)
        return _ensure_non_empty_text(text, "PDF 文件未提取到可用文本内容")

    if ext == ".docx":
        try:
            from docx import Document
        except Exception as exc:
            raise HTTPException(status_code=500, detail=f"DOCX parser unavailable: {exc}") from exc

        doc = Document(io.BytesIO(raw))
        text = "\n".join((p.text or "").strip() for p in doc.paragraphs)
        return _ensure_non_empty_text(text, "DOCX 文件未提取到可用文本内容")

    if ext == ".doc":
        return _extract_doc_text(raw)

    if ext == ".xlsx":
        return _extract_xlsx_text(raw)

    if ext == ".xls":
        return _extract_xls_text(raw)

    raise HTTPException(
        status_code=415,
        detail="不支持的文件类型，仅支持 .txt .md .pdf .doc .docx .xls .xlsx",
    )


def _build_document_list_items(
    db: Session,
    limit: int,
    *,
    source_prefix: str | None = None,
    exclude_prefix: str | None = None,
) -> list[DocumentListItem]:
    query = (
        db.query(
            DocumentTable.doc_id,
            DocumentTable.file_name,
            DocumentTable.version_id,
            DocumentTable.total_lines,
            DocumentTable.created_at,
            DocumentTable.updated_at,
            func.count(ParagraphTable.para_id).label("paragraphs_indexed"),
        )
        .outerjoin(ParagraphTable, ParagraphTable.doc_id == DocumentTable.doc_id)
    )

    if source_prefix:
        query = query.filter(DocumentTable.source_path.like(f"{source_prefix}%"))
    if exclude_prefix:
        query = query.filter(DocumentTable.source_path.notlike(f"{exclude_prefix}%"))

    rows = (
        query.group_by(
            DocumentTable.doc_id,
            DocumentTable.file_name,
            DocumentTable.version_id,
            DocumentTable.total_lines,
            DocumentTable.created_at,
            DocumentTable.updated_at,
        )
        .order_by(DocumentTable.updated_at.desc(), DocumentTable.created_at.desc())
        .limit(limit)
        .all()
    )

    return [
        DocumentListItem(
            doc_id=row.doc_id,
            file_name=row.file_name,
            version_id=row.version_id,
            total_lines=row.total_lines,
            paragraphs_indexed=int(row.paragraphs_indexed or 0),
            created_at=row.created_at,
            updated_at=row.updated_at,
        )
        for row in rows
    ]


def _get_document_list_item(db: Session, doc_id: str) -> DocumentListItem | None:
    row = (
        db.query(
            DocumentTable.doc_id,
            DocumentTable.file_name,
            DocumentTable.version_id,
            DocumentTable.total_lines,
            DocumentTable.created_at,
            DocumentTable.updated_at,
            func.count(ParagraphTable.para_id).label("paragraphs_indexed"),
        )
        .outerjoin(ParagraphTable, ParagraphTable.doc_id == DocumentTable.doc_id)
        .filter(DocumentTable.doc_id == doc_id)
        .group_by(
            DocumentTable.doc_id,
            DocumentTable.file_name,
            DocumentTable.version_id,
            DocumentTable.total_lines,
            DocumentTable.created_at,
            DocumentTable.updated_at,
        )
        .first()
    )
    if not row:
        return None
    return DocumentListItem(
        doc_id=row.doc_id,
        file_name=row.file_name,
        version_id=row.version_id,
        total_lines=row.total_lines,
        paragraphs_indexed=int(row.paragraphs_indexed or 0),
        created_at=row.created_at,
        updated_at=row.updated_at,
    )


# ---------------------------------------------------------------------------
# 5.1 Similarity search endpoint
# ---------------------------------------------------------------------------

@router.post(
    "/search",
    response_model=SearchResponse,
    responses={
        503: {"model": ErrorResponse, "description": "Service not ready"},
        422: {"model": ErrorResponse, "description": "Citation metadata validation failed"},
    },
)
async def similarity_search(
    request: SearchRequest,
    db: Session = Depends(get_session),
):
    """
    Execute a dual-layer legal similarity search.

    Returns ranked document-level matches with nested paragraph-level evidence.
    Each paragraph evidence hit includes mandatory citation metadata
    (file_name, line_start, line_end, version_id) and a match explanation.
    """
    _ensure_retrieval_ready()
    try:
        response = await execute_search(request, db)
        # Validate and enrich results with snippet extraction and version checks
        response.results = validate_and_enrich_results(db, response.results)
        return response
    except TraceabilityValidationError as e:
        return JSONResponse(status_code=422, content=e.error.model_dump())
    except Exception as e:
        logger.exception("Search failed")
        raise HTTPException(status_code=500, detail=str(e))


@router.post(
    "/similar-cases/compare",
    response_model=SimilarCaseSearchResponse,
    responses={
        400: {"model": ErrorResponse, "description": "No uploaded case materials"},
        503: {"model": ErrorResponse, "description": "Service not ready"},
    },
)
async def similar_case_compare(
    request: SimilarCaseSearchRequest,
    db: Session = Depends(get_session),
):
    """Run dedicated similar-case comparison against session-scoped uploads."""
    _ensure_retrieval_ready()
    try:
        return await execute_similar_case_search(request, db)
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Similar-case comparison failed")
        raise HTTPException(status_code=500, detail=str(e))


# ---------------------------------------------------------------------------
# Grounded DeepSeek chat endpoint
# ---------------------------------------------------------------------------

@router.post(
    "/chat",
    response_model=ChatResponse,
    responses={
        503: {"model": ErrorResponse, "description": "Service not ready"},
        422: {"model": ErrorResponse, "description": "Citation metadata validation failed"},
    },
)
async def grounded_chat(
    request: ChatRequest,
    db: Session = Depends(get_session),
):
    """Chat with DeepSeek using only database-grounded legal evidence."""
    _ensure_retrieval_ready()
    try:
        return await execute_grounded_chat(request, db)
    except TraceabilityValidationError as e:
        return JSONResponse(status_code=422, content=e.error.model_dump())
    except Exception as e:
        logger.exception("Grounded chat failed")
        raise HTTPException(status_code=500, detail=str(e))


@router.post(
    "/chat/stream",
    responses={
        503: {"model": ErrorResponse, "description": "Service not ready"},
        422: {"model": ErrorResponse, "description": "Citation metadata validation failed"},
    },
)
async def grounded_chat_stream(
    request: ChatRequest,
    db: Session = Depends(get_session),
):
    """Stream chat output with the same grounding pipeline used by /chat."""
    _ensure_retrieval_ready()
    try:
        return StreamingResponse(
            stream_grounded_chat(request, db),
            media_type="application/x-ndjson",
            headers={
                "Cache-Control": "no-cache",
                "X-Accel-Buffering": "no",
            },
        )
    except TraceabilityValidationError as e:
        return JSONResponse(status_code=422, content=e.error.model_dump())
    except Exception as e:
        logger.exception("Grounded chat stream failed")
        raise HTTPException(status_code=500, detail=str(e))


# ---------------------------------------------------------------------------
# Document ingestion endpoints
# ---------------------------------------------------------------------------

@router.post(
    "/documents",
    response_model=DocumentIngestResponse,
)
async def ingest_document_endpoint(
    request: DocumentIngestRequest,
    db: Session = Depends(get_session),
):
    """Ingest a document into the search index.

    The document will be parsed, segmented into paragraphs, embedded,
    and indexed in both PostgreSQL and Qdrant.
    """
    try:
        result = ingest_document(
            db=db,
            content=request.content,
            file_name=request.file_name,
            source_path=request.source_path,
            version_id=request.version_id,
        )
        return result
    except Exception as e:
        logger.exception("Document ingestion failed")
        raise HTTPException(status_code=500, detail=str(e))


@router.post(
    "/documents/upload",
    response_model=DocumentIngestResponse,
)
async def upload_document(
    file: UploadFile = File(...),
    db: Session = Depends(get_session),
):
    """Upload and ingest a document file.

    Supports .txt .md .pdf .doc .docx .xls .xlsx files.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")

    raw = await file.read()
    content = _extract_upload_text(file.filename, raw)
    result = ingest_document(
        db=db,
        content=content,
        file_name=file.filename,
        source_path=f"upload://{file.filename}",
    )
    return result


@router.get(
    "/documents",
    response_model=list[DocumentListItem],
)
async def list_documents(
    limit: int = 100,
    db: Session = Depends(get_session),
):
    """List uploaded documents persisted in PostgreSQL."""
    limit = max(1, min(limit, 500))
    return _build_document_list_items(
        db=db,
        limit=limit,
        source_prefix=None,
        exclude_prefix="template://",
    )


@router.post(
    "/templates/upload",
    response_model=DocumentListItem,
)
async def upload_template(
    file: UploadFile = File(...),
    db: Session = Depends(get_session),
):
    """Upload and persist a standard contract template."""
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")

    raw = await file.read()
    content = _extract_upload_text(file.filename, raw)
    result = ingest_document(
        db=db,
        content=content,
        file_name=file.filename,
        source_path=f"template://{file.filename}",
    )
    item = _get_document_list_item(db, result.doc_id)
    if not item:
        raise HTTPException(status_code=500, detail="Template upload succeeded but listing record is missing")
    return item


@router.post(
    "/session-files/upload",
    response_model=SessionTempFileItem,
)
async def upload_session_file(
    session_id: str = Form(...),
    kind: SessionTempFileKind = Form(...),
    file: UploadFile = File(...),
):
    """Upload a session-scoped temporary file without persisting it to search storage."""
    session_id = session_id.strip()
    if not session_id:
        raise HTTPException(status_code=400, detail="session_id is required")
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")

    raw = await file.read()
    content = _extract_upload_text(file.filename, raw)
    return session_temp_file_store.add_file(
        session_id=session_id,
        kind=kind,
        file_name=file.filename,
        content=content,
        size_bytes=len(raw),
    )


@router.get(
    "/session-files",
    response_model=list[SessionTempFileItem],
)
async def list_session_files(
    session_id: str,
    kind: SessionTempFileKind | None = None,
):
    """List temporary files for a specific session."""
    session_id = session_id.strip()
    if not session_id:
        raise HTTPException(status_code=400, detail="session_id is required")
    return session_temp_file_store.list_files(session_id=session_id, kind=kind)


@router.delete(
    "/session-files/{file_id}",
    response_model=SessionTempFileItem,
)
async def delete_session_file(file_id: str):
    """Delete a single temporary session file."""
    deleted = session_temp_file_store.delete_file(file_id=file_id)
    if not deleted:
        raise HTTPException(status_code=404, detail=f"Temporary file {file_id} not found")
    return deleted


@router.delete(
    "/session-files/session/{session_id}",
    response_model=SessionTempClearResponse,
)
async def clear_session_files(
    session_id: str,
    kind: SessionTempFileKind | None = None,
):
    """Clear temporary files for a session, optionally filtered by kind."""
    session_id = session_id.strip()
    if not session_id:
        raise HTTPException(status_code=400, detail="session_id is required")
    cleared = session_temp_file_store.clear_session(session_id=session_id, kind=kind)
    return SessionTempClearResponse(session_id=session_id, cleared=cleared, kind=kind)


@router.get(
    "/contract-review/template-recommendation",
    response_model=ReviewTemplateRecommendationResponse,
)
async def get_contract_review_template_recommendation(
    session_id: str,
    db: Session = Depends(get_session),
):
    """Return ranked template recommendations for the current review session."""
    session_id = session_id.strip()
    if not session_id:
        raise HTTPException(status_code=400, detail="session_id is required")

    try:
        return recommend_templates_for_session(session_id=session_id, db=db)
    except Exception as exc:
        logger.exception("Template recommendation failed")
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@router.post(
    "/contract-review/stream",
    responses={
        400: {"model": ErrorResponse, "description": "Invalid request"},
    },
)
async def contract_review_stream(
    request: ContractReviewRequest,
    db: Session = Depends(get_session),
):
    """Stream contract review output against the selected standard template."""
    session_id = request.session_id.strip()
    template_id = request.template_id.strip()
    query = request.query.strip()

    if not session_id:
        raise HTTPException(status_code=400, detail="session_id is required")
    if not template_id:
        raise HTTPException(status_code=400, detail="template_id is required")
    if not query:
        raise HTTPException(status_code=400, detail="query is required")

    return StreamingResponse(
        stream_template_difference_review(
            session_id=session_id,
            template_id=template_id,
            query=query,
            db=db,
        ),
        media_type="application/x-ndjson",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


@router.get(
    "/templates",
    response_model=list[DocumentListItem],
)
async def list_templates(
    limit: int = 100,
    db: Session = Depends(get_session),
):
    """List persisted standard contract templates."""
    limit = max(1, min(limit, 500))
    return _build_document_list_items(
        db=db,
        limit=limit,
        source_prefix="template://",
        exclude_prefix=None,
    )


@router.delete("/templates/{doc_id}")
async def delete_template_endpoint(
    doc_id: str,
    db: Session = Depends(get_session),
):
    """Delete a persisted standard template."""
    template_doc = (
        db.query(DocumentTable.doc_id)
        .filter(DocumentTable.doc_id == doc_id)
        .filter(DocumentTable.source_path.like("template://%"))
        .first()
    )
    if not template_doc:
        raise HTTPException(status_code=404, detail=f"Template {doc_id} not found")

    success = delete_document(db, doc_id)
    if not success:
        raise HTTPException(status_code=404, detail=f"Template {doc_id} not found")
    return {"status": "deleted", "doc_id": doc_id}


@router.delete("/documents/{doc_id}")
async def delete_document_endpoint(
    doc_id: str,
    db: Session = Depends(get_session),
):
    """Delete a document and all associated data."""
    success = delete_document(db, doc_id)
    if not success:
        raise HTTPException(status_code=404, detail=f"Document {doc_id} not found")
    return {"status": "deleted", "doc_id": doc_id}


# ---------------------------------------------------------------------------
# Health and bootstrap status
# ---------------------------------------------------------------------------

@router.get("/health", response_model=BootstrapStatus)
async def health_check():
    """Check the readiness of all system components."""
    from app.core.config import (
        PG_HOST, PG_PORT, PG_USER, PG_PASSWORD, PG_DATABASE,
        QDRANT_HOST, QDRANT_PORT,
    )

    status = BootstrapStatus()

    # Check PostgreSQL
    try:
        import psycopg2
        conn = psycopg2.connect(
            host=PG_HOST, port=PG_PORT,
            user=PG_USER, password=PG_PASSWORD,
            dbname=PG_DATABASE,
        )
        conn.close()
        status.postgresql_ready = True
    except Exception:
        status.postgresql_ready = False

    # Check Qdrant
    try:
        import requests
        resp = requests.get(f"http://{QDRANT_HOST}:{QDRANT_PORT}/healthz", timeout=3)
        status.qdrant_ready = resp.status_code == 200
    except Exception:
        status.qdrant_ready = False

    # Check embedding model (lightweight check - just try to get the model)
    try:
        from app.core.config import EMBEDDING_PROVIDER
        if EMBEDDING_PROVIDER == "google":
            from app.services.embedding import _get_google_client
            _get_google_client()
        else:
            from app.services.embedding import _get_local_model
            _get_local_model()
        status.embedding_model_ready = True
    except Exception:
        status.embedding_model_ready = False

    # Check reranker model
    try:
        from app.services.reranker import _get_reranker
        _get_reranker()
        status.reranker_model_ready = True
    except Exception:
        status.reranker_model_ready = False

    status.all_ready = all([
        status.postgresql_ready,
        status.qdrant_ready,
        status.embedding_model_ready,
        status.reranker_model_ready,
    ])

    if status.all_ready:
        status.message = "All components are ready."
    else:
        missing = []
        if not status.postgresql_ready:
            missing.append("PostgreSQL")
        if not status.qdrant_ready:
            missing.append("Qdrant")
        if not status.embedding_model_ready:
            missing.append("Embedding Model")
        if not status.reranker_model_ready:
            missing.append("Reranker Model")
        status.message = f"Not ready: {', '.join(missing)}"

    return status


@router.post("/bootstrap")
async def trigger_bootstrap():
    """Manually trigger the bootstrap process."""
    from app.core.config import run_bootstrap
    status = run_bootstrap()
    return status
